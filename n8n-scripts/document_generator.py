"""
Document Generator - Create formatted DOCX and PDF from improved resume
Use this in n8n Code node or HTTP Request to a Python service

Required packages:
- python-docx
- reportlab (for PDF) or docx2pdf
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from typing import Dict, Any
import re


def create_resume_docx(
    improved_markdown: str,
    output_path: str,
    style_config: Dict[str, Any] = None
) -> str:
    """
    Create a professionally formatted DOCX from markdown resume

    Args:
        improved_markdown: Resume content in markdown format
        output_path: Path where DOCX should be saved
        style_config: Optional styling configuration

    Returns:
        Path to created DOCX file
    """
    doc = Document()

    # Set up default styles
    if style_config is None:
        style_config = get_default_style_config()

    # Set margins (0.5" top/bottom, 0.75" left/right)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Parse markdown and add to document
    lines = improved_markdown.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # H1 (Name) - Largest, centered
        if line.startswith('# '):
            add_name_header(doc, line[2:], style_config)

        # H2 (Section Headers)
        elif line.startswith('## '):
            add_section_header(doc, line[3:], style_config)

        # H3 (Job titles, etc.)
        elif line.startswith('### '):
            add_job_title(doc, line[4:], style_config)

        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            add_bullet_point(doc, line[2:], style_config)

        # Bold text in markdown **text**
        elif line.startswith('**') and line.endswith('**'):
            add_bold_line(doc, line[2:-2], style_config)

        # Horizontal rule (separator)
        elif line.startswith('---'):
            add_separator(doc)

        # Regular paragraph
        else:
            add_paragraph(doc, line, style_config)

        i += 1

    # Save document
    doc.save(output_path)
    return output_path


def get_default_style_config() -> Dict[str, Any]:
    """Default styling configuration"""
    return {
        'name_font_size': 20,
        'name_color': RGBColor(0, 0, 0),
        'section_header_font_size': 14,
        'section_header_color': RGBColor(31, 78, 121),  # Professional blue
        'job_title_font_size': 11,
        'body_font_size': 10,
        'font_name': 'Calibri',
        'line_spacing': 1.15
    }


def add_name_header(doc: Document, name: str, config: Dict):
    """Add name at the top of resume"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(name)
    run.font.size = Pt(config['name_font_size'])
    run.font.name = config['font_name']
    run.font.bold = True
    run.font.color.rgb = config['name_color']


def add_section_header(doc: Document, header: str, config: Dict):
    """Add section header (EXPERIENCE, EDUCATION, etc.)"""
    # Add small space before section
    doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run(header.upper())
    run.font.size = Pt(config['section_header_font_size'])
    run.font.name = config['font_name']
    run.font.bold = True
    run.font.color.rgb = config['section_header_color']

    # Add bottom border
    para.paragraph_format.border_bottom = True


def add_job_title(doc: Document, title: str, config: Dict):
    """Add job title and company"""
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.size = Pt(config['job_title_font_size'])
    run.font.name = config['font_name']
    run.font.bold = True
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)


def add_bullet_point(doc: Document, text: str, config: Dict):
    """Add bullet point"""
    para = doc.add_paragraph(text, style='List Bullet')
    run = para.runs[0]
    run.font.size = Pt(config['body_font_size'])
    run.font.name = config['font_name']
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(3)


def add_bold_line(doc: Document, text: str, config: Dict):
    """Add bold text line (for dates, locations, etc.)"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(config['body_font_size'])
    run.font.name = config['font_name']
    run.font.bold = True
    para.paragraph_format.space_after = Pt(3)


def add_paragraph(doc: Document, text: str, config: Dict):
    """Add regular paragraph"""
    para = doc.add_paragraph(text)
    run = para.runs[0]
    run.font.size = Pt(config['body_font_size'])
    run.font.name = config['font_name']
    para.paragraph_format.space_after = Pt(6)


def add_separator(doc: Document):
    """Add horizontal line separator"""
    para = doc.add_paragraph()
    para.paragraph_format.border_bottom = True
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> str:
    """
    Convert DOCX to PDF

    Option 1: Using docx2pdf (Windows/Mac)
    Option 2: Using LibreOffice (Linux)
    """
    try:
        # Option 1: docx2pdf
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return pdf_path
    except ImportError:
        # Option 2: LibreOffice command line
        import subprocess
        subprocess.run([
            'libreoffice',
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', os.path.dirname(pdf_path),
            docx_path
        ], check=True)
        return pdf_path


# ===== n8n Usage Example =====
"""
Method 1: Direct in n8n Code Node (Python)
-------------------------------------------
import tempfile
import os

# Get improved markdown from previous node
improved_markdown = $json.improved_markdown

# Create temp files
temp_dir = tempfile.gettempdir()
docx_path = os.path.join(temp_dir, f"resume_{$json.session_id}.docx")
pdf_path = os.path.join(temp_dir, f"resume_{$json.session_id}.pdf")

# Generate DOCX
create_resume_docx(improved_markdown, docx_path)

# Convert to PDF
convert_docx_to_pdf(docx_path, pdf_path)

# Read files as binary
with open(docx_path, 'rb') as f:
    docx_binary = f.read()

with open(pdf_path, 'rb') as f:
    pdf_binary = f.read()

# Return binary data
return [{
    'json': {
        'docx_generated': True,
        'pdf_generated': True
    },
    'binary': {
        'docx': docx_binary,
        'pdf': pdf_binary
    }
}]


Method 2: HTTP Request to Python Microservice
----------------------------------------------
Create a Flask/FastAPI service that receives markdown and returns files.

Then in n8n, use HTTP Request node:
- Method: POST
- URL: http://your-service/generate-resume
- Body: { "markdown": "...", "session_id": "..." }
- Response: URLs to generated files
"""


# ===== Flask Microservice Example =====
"""
from flask import Flask, request, jsonify, send_file
import tempfile
import os

app = Flask(__name__)

@app.route('/generate-resume', methods=['POST'])
def generate_resume():
    data = request.json
    markdown = data.get('improved_markdown')
    session_id = data.get('session_id', 'default')

    # Create files
    temp_dir = tempfile.gettempdir()
    docx_path = os.path.join(temp_dir, f"resume_{session_id}.docx")
    pdf_path = os.path.join(temp_dir, f"resume_{session_id}.pdf")

    # Generate
    create_resume_docx(markdown, docx_path)
    convert_docx_to_pdf(docx_path, pdf_path)

    # Upload to S3 or return as base64
    # ... (implement file upload to S3/storage)

    return jsonify({
        'docx_url': f'https://storage.example.com/resume_{session_id}.docx',
        'pdf_url': f'https://storage.example.com/resume_{session_id}.pdf'
    })

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
"""

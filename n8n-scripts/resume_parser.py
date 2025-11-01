"""
Resume Parser - Extract text from PDF and DOCX files
Use this in n8n Code node (Python)

Required packages in n8n:
- PyPDF2
- python-docx
"""

import io
import json
from typing import Dict, Any

def parse_pdf(file_data: bytes) -> str:
    """Extract text from PDF file"""
    try:
        import PyPDF2

        pdf_file = io.BytesIO(file_data)
        pdf_reader = PyPDF2.PdfReader(pdf_file)

        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"

        return text.strip()
    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")


def parse_docx(file_data: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        from docx import Document

        docx_file = io.BytesIO(file_data)
        doc = Document(docx_file)

        # Extract text from paragraphs
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        return "\n".join(text)
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")


def extract_resume_text(file_data: bytes, mime_type: str) -> Dict[str, Any]:
    """
    Main function to extract text from resume

    Args:
        file_data: Binary file data
        mime_type: MIME type of the file

    Returns:
        Dictionary with extracted text and metadata
    """
    resume_text = ""

    if mime_type == "application/pdf":
        resume_text = parse_pdf(file_data)
        file_type = "pdf"
    elif "wordprocessing" in mime_type or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume_text = parse_docx(file_data)
        file_type = "docx"
    else:
        raise ValueError(f"Unsupported file type: {mime_type}")

    # Basic validation
    if len(resume_text) < 100:
        raise ValueError("Resume text is too short. Please check if the file is valid.")

    # Extract basic metadata
    word_count = len(resume_text.split())
    char_count = len(resume_text)

    return {
        "resume_text": resume_text,
        "file_type": file_type,
        "word_count": word_count,
        "char_count": char_count,
        "extracted_successfully": True
    }


# ===== n8n Code Node Usage =====
"""
In n8n Code Node (Python):

# Get binary data from previous node
file_binary = $binary.data
mime_type = $json.file_type or $json.mimeType

# Parse resume
result = extract_resume_text(file_binary, mime_type)

# Return for next node
return [{
    'json': {
        'resume_text': result['resume_text'],
        'word_count': result['word_count'],
        'char_count': result['char_count'],
        'file_type': result['file_type']
    }
}]
"""
